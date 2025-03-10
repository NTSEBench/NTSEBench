source="2019Delhi"

folderName=f"./inputFiles/{source}/"
image_dir=f"./inputFiles/{source}/images"
docName=folderName+f'{source}.docx'
pdfname=folderName+f'{source}.pdf'

from docxlatex import Document
from PIL import Image
import docx
import os
import fitz  # PyMuPDF
import shutil
import sys
import os
from docx2pdf import convert
import aspose.words as aw


# doc1 = aw.Document(docName)
# if doc1.pages[1].artifacts[1].subtype == aw.Artifact.ArtifactSubtype.WATERMARK:
#         doc1.pages[1].artifacts.delete(doc1.pages[1].artifacts[1])
# doc1.save(pdfname)
# exit(0)
file = open(pdfname, "w")
file.close()
convert(docName, pdfname)

os.makedirs(image_dir)
doc=fitz.open(pdfname)
count=1
for page_num in range(doc.page_count):
        page = doc[page_num]
        image_list = page.get_images()
        for block in page.get_text("dict",sort=True)["blocks"]:
            if(block['type']==1):
                # print("image")
                img_rect = fitz.Rect(block["bbox"])  # Get the bounding box of the image block
                image = page.get_pixmap(clip=img_rect)
                # if(image.width * image.height >= min_image_size):
                pil_image = Image.frombytes("RGB", [image.width, image.height], image.samples)
                pil_image.save(image_dir+"/"+"image"+str(count)+".jpg")
                count+=1

# Create a document
docxfile = Document(docName)
text = docxfile.get_text()
# print(text[:500])
outputDoc = docx.Document()
lines=text.split("\n")
# print(lines)
imgNum=1
imageNames=os.listdir(image_dir)


# for image in imageNames:
    # img=Image.open(image_dir+"/"+image).convert('RGB').save(image_dir+"/"+image)
for line in lines:
    if ("*" in line):
        print("found * in line", line)
    currImageString=f"IMAGE#{imgNum}-image{imgNum}"
    currImageName=f"image{imgNum}"
    if(currImageString in line):
        si=line.index(currImageString)
        if(si!=0):
            # write before part
            p = outputDoc.add_paragraph()
            run = p.add_run(line[0:si])
        # write image
        fullname=[name for name in imageNames if name.startswith(currImageName)][0]
        # print("here")
        # print(fullname)
        outputDoc.add_picture(image_dir+"/"+fullname)
        imgNum+=1
        if(si+len(currImageString)<len(line)):
            # write after part
            p = outputDoc.add_paragraph()
            run = p.add_run(line[(si+len(currImageString)):])
    else:
        #  write into docx
        p = outputDoc.add_paragraph()
        run = p.add_run(line)
        
# Save the document
outputDoc.save(folderName+f"{source}_fixed.docx")
shutil.rmtree(image_dir)

# creating final pdf from docx and saving in parent folder
newInputFile=folderName+f"{source}_fixed.docx"
finalOutputFile=f"./inputFiles/{source}_fixed.pdf"

file = open(finalOutputFile, "w")
file.close()
convert(newInputFile, finalOutputFile)
